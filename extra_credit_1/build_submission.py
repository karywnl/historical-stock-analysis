"""Build the Assignment EC1 practice set report.

Unlike the Assignment 2 builder (which converts the whole analysis notebook into a
docx), EC1 asks for a short 3-5 page report following a template: a cover
table, one short section per exercise with its chart(s) and interpretation,
a critique checklist table, and the filename/submission block. This script
writes that report directly rather than dumping the full EC1 analysis notebook cell
by cell, since the notebook has far more code and detail than fits in five
pages.

Usage: uv run python extra_credit_1/build_submission.py
"""

from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

ASSIGNMENT_ROOT = Path(__file__).resolve().parent
REPO_ROOT = ASSIGNMENT_ROOT.parent
FIGURES_DIR = ASSIGNMENT_ROOT / "figures"
OUTPUT_PATH = ASSIGNMENT_ROOT / "docs" / "submissions" / "26120004_Team2_Assignment-EC1.docx"

REG_NUMBER = "26120004"
TEAM_NUMBER = "2"
PROJECT_TITLE = "Historical Stock Analysis"
TEAM_MEMBERS = "[add team member names here]"
SUBMISSION_DATE = "2026-07-30"

HEADING_COLOR = RGBColor(0x1F, 0x3A, 0x5F)


def add_cover(doc):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Practice Set for Additional Credit (Assignment EC1)")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = HEADING_COLOR

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("Advanced Python Data Visualization")
    sub_run.font.size = Pt(13)
    sub_run.italic = True

    doc.add_paragraph()
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Light Grid Accent 1"
    rows = [
        ("Team Number", TEAM_NUMBER),
        ("Project Title", PROJECT_TITLE),
        ("Team Members", TEAM_MEMBERS),
        ("Date", SUBMISSION_DATE),
    ]
    for i, (label, value) in enumerate(rows):
        table.cell(i, 0).text = label
        table.cell(i, 0).paragraphs[0].runs[0].bold = True
        table.cell(i, 1).text = value

    doc.add_paragraph()
    intro = doc.add_paragraph()
    intro.add_run(
        "This report works through exercises A to F of the EC1 practice set, applied to our own "
        "project (data/stock_data.csv, 49 tickers, 2006-01-02 to 2026-02-20) rather than a textbook "
        "example. It builds on the Assignment 2 and Assignment 3 analysis notebooks, "
        "reusing the same tickers, groupings, and hypotheses. The full working notebook is "
        "extra_credit_1/analysis.ipynb; this document summarizes it in report form."
    )
    doc.add_page_break()


def add_section_heading(doc, text):
    heading = doc.add_heading(level=1)
    run = heading.add_run(text)
    run.font.color.rgb = HEADING_COLOR


def add_body(doc, lines):
    for line in lines:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(line)


def add_image(doc, filename, width_inches=6.0, caption=None):
    path = FIGURES_DIR / filename
    doc.add_picture(str(path), width=Inches(width_inches))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap.add_run(caption)
        cap_run.italic = True
        cap_run.font.size = Pt(9)


def add_image_pair(doc, filename_left, filename_right, caption_left, caption_right, width_inches=2.9):
    table = doc.add_table(rows=2, cols=2)
    table.autofit = True
    left_cell, right_cell = table.rows[0].cells
    left_cell.paragraphs[0].add_run().add_picture(str(FIGURES_DIR / filename_left), width=Inches(width_inches))
    right_cell.paragraphs[0].add_run().add_picture(str(FIGURES_DIR / filename_right), width=Inches(width_inches))
    for cell, caption in ((table.rows[1].cells[0], caption_left), (table.rows[1].cells[1], caption_right)):
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(caption)
        run.italic = True
        run.font.size = Pt(9)


def add_exercise_a(doc):
    add_section_heading(doc, "Exercise A: Redesign for Clarity")
    doc.add_paragraph(
        "Chosen chart: Assignment 2 Figure 8, six tickers (NVDA, AVGO, TSM, ASML, AAPL, MSFT) "
        "indexed to 100 since 2009-08-06."
    )
    add_body(doc, [
        "Wrong scale for the data's range: NVDA ends 27 times above MSFT, so a linear y-axis "
        "leaves four of six tickers flat near zero for most of the chart.",
        "Poor pre-attentive encoding of change: a linear axis encodes equal absolute change, not "
        "equal percent change, which is the wrong read for an indexed growth chart.",
        "Label collision: ASML, TSM, and AAPL end within 15% of each other, so their direct end "
        "labels overlap.",
    ])
    add_image_pair(doc, "ec1_a_before.png", "ec1_a_after.png",
                   "Before: linear y-axis", "After: log y-axis, nudged labels")
    add_body(doc, [
        "Switching to a log y-axis makes equal vertical space mean equal percent change, so all "
        "six tickers show a visible growth curve across the full 16 years, not just the two winners.",
        "Nudging the three closest end labels apart by a small fixed amount keeps all six names "
        "readable instead of stacking three of them illegibly.",
    ])


def add_exercise_b(doc):
    add_section_heading(doc, "Exercise B: Small Multiples Storytelling")
    doc.add_paragraph(
        "Categorical variable: region (US, Korea, Hong Kong, Switzerland, Paris; same one ticker "
        "per region as Assignment 3). Relationship shown: trading volume vs. each ticker's own "
        "2015-2017 calm baseline, across the four periods, one subplot per region on a shared y-axis."
    )
    add_image(doc, "ec1_b_small_multiples.png", width_inches=5.5)
    add_body(doc, [
        "Assignment 3's combined chart averaged all five regions into one 2008 bar (about 2.7x "
        "normal volume), which looked like a broad, even panic.",
        "Split by region, the US panel (JPM) spikes to about 5.1x normal in 2008, clearly the "
        "highest bar in the whole figure, while the other four regions sit at 1.5-2.6x, closer to "
        "their own 2020 numbers. Averaging JPM in with four calmer regions hid how much of an "
        "outlier the US region was that period, matching Assignment 3's own caveat about JPM's "
        "Bear Stearns acquisition news.",
    ])


def add_exercise_c(doc):
    add_section_heading(doc, "Exercise C: Interactive Linked Views")
    doc.add_paragraph(
        "Two linked Altair views on the H1 tickers (chip/AI vs. rest of tech, since 2023-01-03): "
        "a scatter of daily return over time, and a bar chart of average daily return per ticker. "
        "A selection_interval brush on the scatter's date axis filters the bar chart via "
        "transform_filter; a bind_radio dropdown isolates either group."
    )
    add_image_pair(doc, "ec1_c_linked_views.png", "ec1_c_linked_views_brushed.png",
                   "Default (unfiltered) state", "After dragging a brush over Jan-Aug 2025")
    add_body(doc, [
        "A static chart only answers one question: which group did better over the whole period "
        "(chip/AI names take 4 of the top 5 spots). The brush turns that into 'did the chip group "
        "also win during this specific window,' recomputed instantly instead of redrawn by hand.",
        "The group dropdown answers whether one group's result is carried by a single name or "
        "spread across it, by hiding the other group instead of asking the reader to filter by color.",
        "The brushed screenshot shows the effect: over just Jan-Aug 2025 the ranking reshuffles, "
        "ORCL and AMD move to the top and AAPL turns negative, which the full-period bar chart alone "
        "does not show.",
    ])


def add_exercise_d(doc):
    add_section_heading(doc, "Exercise D: Statistical Rigor")
    doc.add_paragraph(
        "Hypothesis: H1, the chip/AI rally is narrower than the rest of tech. Mean daily return "
        "per group since 2023-01-03, with 95% confidence intervals, plus a Welch's t-test."
    )
    add_image(doc, "ec1_d_h1_stats.png", width_inches=3.4)
    add_body(doc, [
        "Chip/AI group: 0.26% mean daily return, 95% CI [0.19%, 0.34%]. Rest of tech: 0.14%, "
        "95% CI [0.09%, 0.19%]. Welch's t-test: t = 2.69, p = 0.0073.",
        "Verdict: this supports H1. The gap survives a test of the average ordinary day, not just "
        "a handful of large days compounding in the indexed price chart, and p = 0.0073 is well "
        "under the usual 0.05 cutoff.",
        "Honest complication: the daily gap is small and the two groups' confidence intervals "
        "almost touch. It only becomes the dramatic split seen in Assignment 3's Figure A because "
        "it compounds over three years of trading days.",
    ])


def add_exercise_e(doc):
    add_section_heading(doc, "Exercise E: Reusable Style Module")
    doc.add_paragraph(
        "Style module: extra_credit_1/chart_style.py. A fixed six-color Okabe-Ito palette (validated with "
        "the course's colorblind-safety checker), the same two colors for the chip/AI and rest of "
        "tech groups everywhere they appear, standard font sizes, and a standard figure size."
    )
    add_image_pair(doc, "ec1_e_figure_a_restyled.png", "ec1_e_figure_b_restyled.png",
                   "Figure A, restyled", "Figure B, restyled")
    add_body(doc, [
        "The palette passed the colorblind separation checks for every pair used, rather than "
        "being picked by eye.",
        "Blue and orange are reserved for the same two groups everywhere in the project, so color "
        "follows the entity instead of being re-picked per chart.",
        "These are categorical groups, not a magnitude, so a categorical palette is the right "
        "encoding rather than a single-hue gradient.",
    ])


def add_exercise_f(doc):
    add_section_heading(doc, "Exercise F: Widening the Region Basket (H2/H3 Follow-Up)")
    doc.add_paragraph(
        "Assignment 3 flagged widening the region basket from one stock to two or three as its "
        "next step, since JPM's own Bear Stearns news looked like it was drowning out the regional "
        "signal in 2008. Widened baskets, keeping only full-history tickers: US (JPM, AAPL, XOM), "
        "Korea (005930.KS, 000660.KS), Hong Kong (0700.HK, 0939.HK, 1398.HK), Switzerland (NOVN.SW, "
        "ROG.SW), Paris (MC.PA only, the sole Paris-listed ticker in the dataset)."
    )
    add_image(doc, "ec1_f_widened_basket.png", width_inches=5.5)
    add_body(doc, [
        "One stock per region: 2008 correlation (0.250) was almost identical to the calm baseline "
        "(0.250). Two to three stocks per region: 2008 moves to 0.326 vs. a calm baseline of 0.282, "
        "now clearly above baseline instead of sitting on top of it.",
        "2020 and 2022 barely move under the widened basket, so the single-stock approach was not "
        "wrong everywhere, it specifically undersold 2008, the one period where JPM had its own "
        "large company-specific news on top of the broader crisis.",
        "Limitation: Korea and Switzerland baskets are still sector-narrow (tech, pharma), and "
        "Paris cannot be widened at all with this dataset, so 2008 for those regions still rests "
        "on thinner ground than the US and Hong Kong baskets.",
    ])
    doc.add_paragraph()


def add_critique_checklist(doc):
    add_section_heading(doc, "Peer Critique Checklist")
    doc.add_paragraph(
        "Self-review of Figure F (Exercise F), completed ahead of the peer exchange:"
    )
    table = doc.add_table(rows=6, cols=3)
    table.style = "Light Grid Accent 1"
    header = ["Question", "Rating / Answer", "One Improvement"]
    for i, text in enumerate(header):
        table.cell(0, i).text = text
        table.cell(0, i).paragraphs[0].runs[0].bold = True

    rows = [
        ("Is the chart type appropriate for this data and question?", "Yes, 5/5",
         "None needed; grouped bars fit comparing two methods across four discrete periods."),
        ("Is there unnecessary chart junk (gridlines, borders, 3D, redundant color)?", "No, 5/5",
         "Period labels could carry a shared date subtitle instead of repeating years."),
        ("Are axes, labels, and legends clear without needing explanation?", "Yes, 4/5",
         "Y-axis label should say correlation of daily returns, not just 'correlation'."),
        ("Is color used meaningfully, not just decoratively?", "Yes, 5/5",
         "None needed; blue/orange separate the two methods, consistent with every other figure."),
        ("Does the written interpretation match what the chart actually shows?", "Yes, 5/5",
         "None needed; the write-up also reports that 2020 and 2022 barely moved."),
    ]
    for i, (question, rating, improvement) in enumerate(rows, start=1):
        table.cell(i, 0).text = question
        table.cell(i, 1).text = rating
        table.cell(i, 2).text = improvement

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.add_run(
        "The peer half of this exercise, exchanging Figure F with another team and receiving their "
        "checklist in return, is left for that in-person/async exchange and is not filled in here."
    ).italic = True


def add_filename_block(doc):
    add_section_heading(doc, "Filename & Submission")
    table = doc.add_table(rows=3, cols=1)
    table.cell(0, 0).text = f"Filename: {REG_NUMBER}_Team{TEAM_NUMBER}_Assignment-EC1.pdf"
    table.cell(1, 0).text = "Submit via: Digiicampus"
    table.cell(2, 0).text = "Due: 31st July, 2026 11:59 PM"


def main():
    doc = Document()
    add_cover(doc)
    add_exercise_a(doc)
    add_exercise_b(doc)
    add_exercise_c(doc)
    add_exercise_d(doc)
    add_exercise_e(doc)
    add_exercise_f(doc)
    add_critique_checklist(doc)
    add_filename_block(doc)
    doc.save(OUTPUT_PATH)
    print(f"wrote {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
