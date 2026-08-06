"""Build the Assignment 4 report (R for Data Visualization: Fundamentals & First EDA).

Follows the same pattern as build_ec1_docx.py: a cover table matching the
assignment's own template fields, then Part A (setup evidence, pasted as
console text rather than screenshots, which the worksheet explicitly allows),
Part B (left as a placeholder for the teammate handling the concept check),
and Part C (the full R EDA, mirroring eda.ipynb's structure and figures).

The R code itself lives in assignment4_eda.R at the repo root; this script
only assembles the written report around it.

Usage: uv run python scripts/build_assignment4_docx.py
"""

from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

REPO_ROOT = Path(__file__).resolve().parent.parent
ASSIGNMENTS_DIR = REPO_ROOT / "assignments"
FIGURES_DIR = ASSIGNMENTS_DIR / "figures" / "assignment4"
OUTPUT_PATH = ASSIGNMENTS_DIR / "26120004_Team2_Assignment4.docx"

REG_NUMBER = "26120004"
TEAM_NUMBER = "2"
PROJECT_TITLE = "Historical Stock Analysis"
DATASET_LINE = (
    "https://www.kaggle.com/datasets/ibrahimshahrukh/top-50-companies-dataset "
    "(note: this Kaggle dataset has since been removed/taken down by its uploader)"
)
TEAM_MEMBERS = "Karthikeyan M (26120004), Mirthula M (26120121), Lakshmi Mayuri Kavya N (26120161)"
SUBMISSION_DATE = "[team: confirm submission date]"

HEADING_COLOR = RGBColor(0x1F, 0x3A, 0x5F)


def add_cover(doc):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Assignment 4")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = HEADING_COLOR

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("R for Data Visualization: Fundamentals & First EDA")
    sub_run.font.size = Pt(13)
    sub_run.italic = True

    doc.add_paragraph()
    table = doc.add_table(rows=5, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Light Grid Accent 1"
    rows = [
        ("Team Number", TEAM_NUMBER),
        ("Project Title", PROJECT_TITLE),
        ("Dataset (name & source/link)", f"data/stock_data.csv - {DATASET_LINE}"),
        ("Team Members", TEAM_MEMBERS),
        ("Date", SUBMISSION_DATE),
    ]
    for i, (label, value) in enumerate(rows):
        table.cell(i, 0).text = label
        table.cell(i, 0).paragraphs[0].runs[0].bold = True
        table.cell(i, 1).text = value

    doc.add_paragraph()
    intro = doc.add_paragraph()
    intro.add_run(
        "This report redoes our Assignment 2 python eda (eda.ipynb) in r, same dataset, same "
        "questions, a different language, per the assignment's own instructions. Part C figure "
        "numbers below are r figures 1 through 8, separate from the python eda's own figure "
        "numbering. The full r script is assignment4_eda.R at the repo root."
    )
    doc.add_page_break()


def add_section_heading(doc, text, level=1):
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    run.font.color.rgb = HEADING_COLOR


def add_body(doc, lines):
    for i, line in enumerate(lines):
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(line)
        if i < len(lines) - 1:
            doc.add_paragraph()


def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def add_image(doc, filename, width_inches=6.0, caption=None):
    path = FIGURES_DIR / filename
    doc.add_picture(str(path), width=Inches(width_inches))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap.add_run(caption)
        cap_run.italic = True
        cap_run.font.size = Pt(9)


# ---------------------------------------------------------------------------
# part a

def add_part_a(doc):
    add_section_heading(doc, "Part A: Installation & Setup")
    doc.add_paragraph(
        "R and RStudio were installed via Homebrew (brew install --cask r rstudio), which is "
        "equivalent to the CRAN/Posit installers. Console output is pasted below in place of "
        "screenshots, which the worksheet explicitly allows (\"paste screenshots or console "
        "output\")."
    )

    doc.add_paragraph().add_run("Packages installed and loaded:").bold = True
    add_code_block(doc, (
        "> library(tidyverse)\n"
        "-- Attaching core tidyverse packages -------------------- tidyverse 2.0.0 --\n"
        "v dplyr     1.2.1     v readr     2.2.0\n"
        "v forcats   1.0.1     v stringr   1.6.0\n"
        "v ggplot2   4.0.3     v tibble    3.3.1\n"
        "v lubridate 1.9.5     v tidyr     1.3.2\n"
        "v purrr     1.2.2\n"
        "-- Conflicts ------------------------------------- tidyverse_conflicts() --\n"
        "x dplyr::filter() masks stats::filter()\n"
        "x dplyr::lag()    masks stats::lag()\n"
        "> library(data.table)\n"
        "data.table 1.18.4 using 5 threads (see ?getDTthreads).\n"
        "> library(lubridate)\n"
        "> library(reshape2)\n"
        "(all four load cleanly, no error text - the lines above are the standard\n"
        "\"masking\" notices every tidyverse session prints, not failures)"
    ))

    doc.add_paragraph().add_run("Dataset loaded and checked:").bold = True
    add_code_block(doc, (
        "> df <- read.csv('data/stock_data.csv')\n"
        "> dim(df)\n"
        "[1] 230111      8\n"
        "> names(df)\n"
        "[1] \"Date\"      \"Ticker\"    \"Open\"      \"High\"      \"Low\"       \"Close\"\n"
        "[7] \"Adj.Close\" \"Volume\"\n"
        "> str(df)\n"
        "'data.frame':   230111 obs. of  8 variables:\n"
        " $ Date     : chr  \"2006-01-02\" \"2006-01-03\" \"2006-01-04\" ...\n"
        " $ Ticker   : chr  \"000660.KS\" \"000660.KS\" \"000660.KS\" ...\n"
        " $ Open     : num  36100 38100 38250 35900 34950 ...\n"
        " $ High     : num  37650 38400 39050 36000 35850 ...\n"
        " $ Low      : num  35850 36850 35300 33600 34850 ...\n"
        " $ Close    : num  37600 38200 35400 34750 35100 ...\n"
        " $ Adj.Close: num  27358 27795 25758 25285 25539 ...\n"
        " $ Volume   : num  13334311 13815168 23064896 21043236 12557226 ...\n"
        "> head(df)\n"
        "        Date    Ticker  Open  High   Low Close Adj.Close   Volume\n"
        "1 2006-01-02 000660.KS 36100 37650 35850 37600  27358.40 13334311\n"
        "2 2006-01-03 000660.KS 38100 38400 36850 38200  27794.97 13815168\n"
        "3 2006-01-04 000660.KS 38250 39050 35300 35400  25757.65 23064896\n"
        "4 2006-01-05 000660.KS 35900 36000 33600 34750  25284.69 21043236\n"
        "5 2006-01-06 000660.KS 34950 35850 34850 35100  25539.36 12557226\n"
        "6 2006-01-09 000660.KS 35600 36750 35450 36750  26739.93 13524015"
    ))

    add_body(doc, [
        "230,111 rows and 8 columns, matching what we expected and what the python eda found. "
        "R renamed Adj Close to Adj.Close automatically, since read.csv does not allow spaces in "
        "column names, and Date loaded as plain text rather than a date type, both just how r's "
        "base reader works rather than anything wrong with the data.",
    ])
    doc.add_page_break()


# ---------------------------------------------------------------------------
# part b

def add_qa(doc, question, answer_lines, note=None):
    q = doc.add_paragraph()
    q.add_run(question).bold = True
    for line in answer_lines:
        doc.add_paragraph(line, style="List Bullet")
    if note:
        n = doc.add_paragraph()
        n.add_run(note).italic = True
    doc.add_paragraph()


def add_part_b(doc):
    add_section_heading(doc, "Part B: Concept Check")

    add_qa(doc,
        "B.1 In your own words, distinguish descriptive statistics from inferential "
        "statistics. Give one example of each using your own project dataset.",
        [
            "Descriptive statistics summarize and describe the main characteristics of a "
            "dataset using measures such as the mean, median, standard deviation, and "
            "graphs. Inferential statistics use sample data to make conclusions about a "
            "larger population.",
            "Descriptive example: in our stock market dataset, we calculated the mean "
            "daily return and median trading volume to summarize the data.",
            "Inferential example: we performed a hypothesis test to check whether the "
            "average daily returns of two markets were significantly different.",
        ])

    add_qa(doc,
        "B.2 Suppose your key numeric variable has a few extreme outliers. Which measure "
        "of central tendency would you trust more - the mean or the median - and why? "
        "Name one measure of spread that is similarly robust to outliers.",
        [
            "We would choose the median, since it is resistant to outliers, whereas the "
            "mean is heavily affected by extreme values.",
        ],
        note="team to add: a robust measure of spread (e.g. the interquartile range, IQR) "
             "still needs to be named to fully answer this question.")

    add_qa(doc,
        "B.3 What two parameters define a normal distribution, and what does each one "
        "control?",
        [
            "A normal distribution is defined by two parameters: the mean, which "
            "determines the central location of the distribution, and the standard "
            "deviation, which controls the spread of the data away from the mean, "
            "expressed in the original data's units.",
        ])

    add_qa(doc,
        "B.4 Name one sampling method other than simple random sampling, and describe a "
        "situation from your project domain where it would be the better choice.",
        [
            "One sampling method is stratified sampling, where the population is divided "
            "into homogeneous groups and samples are drawn from each group.",
            "In our stock market dataset, the data contains companies from different "
            "countries, so we would divide the data into groups based on country and "
            "randomly sample from each group. This ensures every country is properly "
            "represented.",
        ])

    add_qa(doc,
        "B.5 Explain, in a sentence or two, what a p-value below your chosen significance "
        "level (alpha) tells you about the null hypothesis.",
        [
            "A p-value below the chosen significance level (alpha) means there is enough "
            "evidence to reject the null hypothesis. It suggests that the observed result "
            "is statistically significant and is unlikely to have happened by chance.",
        ])

    add_qa(doc,
        "B.6 What is the difference between base R, the tidyverse, and data.table as "
        "approaches to working with data? Name one advantage of each.",
        [
            "Base R uses built-in R functions for data analysis. Advantage: no extra "
            "packages are required.",
            "tidyverse is a collection of packages that make data manipulation and "
            "visualization simple and readable. Advantage: easy-to-read syntax and an "
            "efficient workflow.",
            "data.table is a package designed for fast data manipulation, especially with "
            "large datasets. Advantage: very fast and memory-efficient.",
        ])

    add_qa(doc,
        "B.7 Which single dplyr verb would you use to (a) keep only the rows meeting a "
        "condition, and (b) collapse many rows into one summary row per group?",
        [
            "(a) To keep only rows that meet a condition, use filter().",
            "(b) To create one summary row per group, use summarise(), paired with "
            "group_by() to define the groups.",
        ])

    add_qa(doc,
        "B.8 Explain the difference between wide and long data format. Name one function "
        "- from reshape2 or tidyr - that converts data from wide to long.",
        [
            "Wide format stores each variable in a separate column. Long format stores "
            "values from multiple columns in a single column, with another column "
            "indicating the variable name.",
            "One function that converts data from wide to long is pivot_longer() from "
            "tidyr.",
        ])

    add_qa(doc,
        "B.9 List the seven layers of ggplot2's Grammar of Graphics, in order.",
        [
            "Data",
            "Aesthetics (aes)",
            "Geometries (geom)",
            "Statistics (stat)",
            "Scales",
            "Coordinate system (coord)",
            "Facets",
        ])

    add_qa(doc,
        "B.10 Your scatter plot has hundreds of points overlapping in one corner. Name "
        "two techniques you could use in ggplot2 to make the underlying pattern more "
        "visible.",
        [
            "Use transparency by setting the alpha value, so overlapping points become "
            "visible rather than solid.",
            "Use geom_jitter() to slightly spread overlapping points and reveal the "
            "underlying pattern.",
        ])

    doc.add_page_break()


# ---------------------------------------------------------------------------
# part c0

def add_part_c0(doc):
    add_section_heading(doc, "Part C: First EDA in R")
    add_section_heading(doc, "C0: Setup & Data Overview", level=2)

    doc.add_paragraph().add_run("C0.1 shape and column types").bold = True
    add_body(doc, [
        "The dataset has 230,111 rows and 8 columns: Date (character), Ticker (character), "
        "Open, High, Low, Close, Adj.Close, and Volume (all numeric). This matches what we "
        "expected, one row per ticker per trading day across 49 tickers, no surprises in row "
        "count, and it matches the python eda's own dtypes exactly aside from R's automatic "
        "Adj Close -> Adj.Close renaming and Date loading as text instead of a date type.",
    ])

    doc.add_paragraph().add_run("C0.2 summary statistics").bold = True
    add_body(doc, [
        "Volume has the widest range by far, 0 to about 18.8 billion, but that is a share "
        "count, not a price, so it is not really comparable to the price columns.",
        "Among the price columns, High has the widest range, about 955,000, stretched by "
        "Korean won prices sitting at a completely different scale from dollar prices.",
        "Volume also has the largest mean minus median gap, about 44 million, and every price "
        "column shows the same roughly 40x mean-to-median gap for the same reason: mixed "
        "currencies pulling the average up, not genuine skew in any single stock. All of this "
        "matches the python eda's C0.2 numbers exactly.",
    ])

    doc.add_paragraph().add_run("C0.3 missing data").bold = True
    add_body(doc, [
        "colSums(is.na(df)) shows zero missing values in every column, identical to the python "
        "eda. Nothing needs to be dropped or imputed for missingness; the real quirks in this "
        "data (mixed currencies, staggered start dates) do not show up as missing values.",
    ])
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# part c1

def add_part_c1(doc):
    add_section_heading(doc, "C1: Univariate Exploration", level=2)

    doc.add_paragraph().add_run("C1.1 distribution of the key numeric variable").bold = True
    add_image(doc, "fig1_daily_return_hist.png", width_inches=5.5)
    add_body(doc, [
        "Daily return is the numeric variable most central to the project, the same choice as "
        "the python eda, since raw price cannot be compared across the 49 tickers once currency "
        "is mixed in.",
        "Figure 1 is roughly symmetric and centered near zero, mean about 0.08 percent and "
        "median about 0.04 percent a day, which makes sense since most days a stock does not "
        "move much either way.",
        "The shape is not a clean bell curve: a sharp narrow peak right at zero, with long thin "
        "tails stretching to about minus 35 percent and plus 52 percent in a single day. Those "
        "extreme days are rare but far more common than a normal distribution would predict, "
        "the same read the python eda reached.",
    ])

    doc.add_paragraph().add_run("C1.2 category counts for a key categorical variable").bold = True
    add_image(doc, "fig2_market_counts.png", width_inches=5.5)
    add_body(doc, [
        "Market, derived from ticker suffix, is the categorical variable again. Figure 2 shows "
        "the same heavy skew as the python eda: United States rows are 184,793 of 230,111 "
        "total, about 80 percent, while Saudi Arabia has only 1,550 and Paris only 5,152.",
        "This is mostly a ticker-count imbalance (39 of 49 tickers are US-listed), not a market "
        "trading more often, which means the Paris and Saudi Arabia labels are really single "
        "company statistics.",
    ])

    doc.add_paragraph().add_run("C1.3 box plot for outliers (optional / bonus)").bold = True
    add_image(doc, "fig3_daily_return_box.png", width_inches=5.5)
    add_body(doc, [
        "Figure 3's box sits tightly between about minus 3.5 percent and plus 3.7 percent, with "
        "a dense scatter of points beyond both whiskers, about 7.24 percent of all ticker-days, "
        "identical to the python eda's outlier fence and share.",
        "The most extreme days on record are the same real events found in python: Netflix's "
        "two roughly 35 percent single-day drops (2022-04-20 and 2011-10-25), Nvidia's 31 "
        "percent drop during the 2008 crisis, Bank of America's swings in 2008-2009, Meta's 26 "
        "percent drop in 2022, and AMD's 52 percent jump in 2016. These read as genuine extreme "
        "trading days tied to real news, not data errors, so they should stay in the dataset.",
    ])
    doc.add_page_break()


# ---------------------------------------------------------------------------
# part c2

def add_part_c2(doc):
    add_section_heading(doc, "C2: Relationships Between Variables", level=2)

    doc.add_paragraph().add_run("C2.1 scatter plot of two related numeric variables").bold = True
    add_image(doc, "fig4_volume_vs_move.png", width_inches=5.5)
    add_body(doc, [
        "Same question as the python eda: does bigger trading volume come with bigger price "
        "moves. The Pearson correlation between volume and the size of the daily move is 0.207, "
        "identical to the python result.",
        "Figure 4 shows a widening funnel, not a tight line: low volume days are almost always "
        "small moves, but high volume days can be small or huge. Volume raises the ceiling on "
        "how big a move can get without guaranteeing a big move happens, a weak relationship at "
        "best, the same conclusion the python eda reached.",
    ])

    doc.add_paragraph().add_run("C2.2 correlation heatmap").bold = True
    add_image(doc, "fig5_correlation_heatmap.png", width_inches=5.0)
    add_body(doc, [
        "Same eight tickers as the python eda: five chip names (NVDA, AVGO, TSM, ASML, AMD), "
        "Apple, Microsoft, and one non-US name outside tech, LVMH (MC.PA).",
        "Figure 5 shows the same clustering: ASML-TSM is the strongest pair at 0.63, and the "
        "five chip names sit together at 0.53-0.63, noticeably tighter than their correlation "
        "with Apple or Microsoft (0.38-0.54).",
        "LVMH is the weakest-correlated name against everything else, 0.25-0.40, well below any "
        "tech-to-tech pair. Same caveat as the python eda: this cannot separate sector from "
        "region since LVMH is both non-tech and the only non-US name here, and it is a full "
        "history average, not a check on crisis periods specifically.",
    ])

    doc.add_paragraph().add_run(
        "C2.3 recreate one plot above using a second R graphics system"
    ).bold = True
    add_image(doc, "fig6_volume_vs_move_lattice.png", width_inches=5.5)
    add_body(doc, [
        "Figure 6 recreates Figure 4 (volume vs. size of move) using Lattice's xyplot(), a "
        "different graphics system from the ggplot2 used above, on the same sampled 3,000 "
        "points.",
        "The shape reads the same in both. Lattice's default styling is plainer and its log "
        "axis prints as powers of ten (10^2, 10^4) rather than ggplot2's 1e+03 style, a small "
        "but real difference in how the two systems format the same information. ggplot2 was "
        "quicker to get labeled and styled the way we wanted; Lattice needed less code for a "
        "quick default plot but less control over the details.",
    ])
    doc.add_page_break()


# ---------------------------------------------------------------------------
# part c3

def add_part_c3(doc):
    add_section_heading(doc, "C3: Group & Categorical Comparisons", level=2)

    doc.add_paragraph().add_run(
        "C3.1 group by a key categorical variable, compare a numeric outcome"
    ).bold = True
    add_image(doc, "fig7_mean_return_by_market.png", width_inches=5.5)
    add_body(doc, [
        "Grouping by market again, Figure 7 shows the same order the python eda found: Korea "
        "has the highest mean daily return (0.094 percent), then United States (0.086 percent), "
        "Hong Kong, Paris, Switzerland, and Saudi Arabia lowest (0.015 percent).",
        "Korea is also the most volatile market by standard deviation (2.37 percent daily) and "
        "Saudi Arabia the calmest (1.05 percent), so higher average return and higher "
        "volatility travel together here, same as the python eda found.",
        "Same caveat carried over from C1.2: Saudi Arabia, Paris, and Switzerland are one or "
        "two tickers each, so this is not really a market-wide statistic, it is a statistic "
        "about a handful of individual companies wearing a market label.",
    ])

    doc.add_paragraph().add_run("C3.2 key variable over time").bold = True
    add_image(doc, "fig8_indexed_price.png", width_inches=6.0)
    add_body(doc, [
        "Figure 8 indexes NVDA, AVGO, TSM, ASML, AAPL, and MSFT to 100 at 2009-08-06, "
        "Broadcom's start date and the latest starter of the six, the same rebase logic as the "
        "python eda.",
        "NVDA reaches about 63,100 by 2026-02-20 and AVGO about 29,200, both far ahead of ASML "
        "and TSM (about 58-62x) and AAPL (about 54x), with MSFT trailing at about 23x, matching "
        "the python eda's Figure 8 numbers.",
        "This supports the AI rally narrowness idea but complicates the simple version of it: "
        "AVGO alone would still justify a chip-outperformance story without NVDA, but ASML and "
        "TSM end up closer to AAPL's growth than to NVDA or AVGO's.",
    ])
    doc.add_page_break()


# ---------------------------------------------------------------------------
# part c4 / c5

def add_part_c4(doc):
    add_section_heading(doc, "C4: Data Quality Notes", level=2)
    add_body(doc, [
        "Mixed currencies across tickers - revealed in C0.2's summary table, the roughly 40x "
        "gap between mean and median in every price column is currency mixing, not real skew, "
        "since Korean won prices sit at a totally different scale from dollar prices. Every "
        "chart that combines tickers in this script works around it with returns or indexing "
        "instead of raw price.",
        "Heavy market imbalance - revealed in Figure 2, United States is about 80 percent of "
        "all rows, and two of six market labels (Paris, Saudi Arabia) are backed by a single "
        "company each, so any market-level statistic in Figure 7 is really a statement about "
        "one or two companies.",
        "A small number of zero-volume days concentrated in one ticker - this one came out of "
        "building Figure 4/6 in R specifically. R's scale_x_log10() threw a warning about "
        "infinite values from the log transform, which led to checking sum(df$Volume == 0). "
        "1,706 rows (0.74 percent of the dataset) have Volume recorded as exactly zero, and "
        "1,459 of those, about 85 percent, belong to a single ticker, AZN, spread mostly across "
        "2009-2019 with a smaller cluster in 2024-2025. That is not spread evenly across all 49 "
        "tickers, so it looks like a data recording gap specific to AZN rather than genuinely "
        "zero trading on those days for a large pharmaceutical stock.",
    ])
    doc.add_paragraph()


def add_part_c5(doc):
    add_section_heading(doc, "C5: Synthesis & Reflection", level=2)

    doc.add_paragraph().add_run("C5.1 three key findings").bold = True
    add_body(doc, [
        "Figure 8 shows the AI rally narrowness idea is real but uneven within the chip group: "
        "NVDA and AVGO pulled dramatically ahead (631x and 292x their 2009 value) while ASML "
        "and TSM only reached about 58-62x, barely ahead of AAPL's 54x. The finding should "
        "probably be about specific standout names rather than the whole chip sector as a "
        "block.",
        "Figure 5 shows the five chip names correlate more tightly with each other (0.53-0.63) "
        "than with AAPL or MSFT (0.38-0.54), and far more than with LVMH (0.25-0.40). That is "
        "independent evidence for the same narrowness story from a different angle, not just "
        "growth, but day-to-day correlation too.",
        "Figures 1 and 3 together show daily returns sharply peaked near zero with fat tails, "
        "about 7.2 percent of ticker-days sitting outside the box plot fence, and the most "
        "extreme days lining up with specific known events like the 2008 crisis or company "
        "earnings shocks rather than being spread evenly through time.",
    ])

    doc.add_paragraph().add_run(
        "C5.2 which R graphics system felt most natural"
    ).bold = True
    add_body(doc, [
        "ggplot2 felt the most natural for this dataset. Its grammar-of-graphics approach, "
        "mapping columns directly to x, y, color, and building up labs() and geoms as layers, "
        "was the closest match to how we already think about a chart after doing the python eda "
        "in matplotlib and seaborn.",
        "Base R's barplot() (Figure 2) was fine for one simple bar chart but needed more manual "
        "margin adjustment (par(mar = ...)) than ggplot2 to stop axis labels from clipping.",
        "Lattice (Figure 6) worked well for a quick recreation but felt the least flexible of "
        "the three beyond its built-in plot types; its default styling was plainer and harder "
        "to customize to match the rest of the figures.",
    ])

    doc.add_paragraph().add_run(
        "C5.3 comparison to the original python EDA"
    ).bold = True
    add_body(doc, [
        "Almost nothing looked different. Every number recomputed in R matched the python eda "
        "to the same precision shown: the same 0.207 volume-return correlation, the same 0.63 "
        "ASML-TSM correlation, the same 7.24 percent outlier share, the same rebase date and "
        "indexed values in Figure 8. That is reassuring rather than surprising, since both "
        "notebooks read the same csv file and just perform the same operations in a different "
        "language.",
        "One genuine difference: R's ggplot2 explicitly warned about the log-10 transform "
        "producing infinite values while building Figure 4, which led directly to finding the "
        "1,706 zero-volume rows concentrated in AZN (see C4). The python matplotlib version of "
        "the same chart did not surface an equivalent warning in the notebook output, so this "
        "specific data quality issue was easier to catch in R.",
        "A smaller mechanical difference: R renamed Adj Close to Adj.Close on load, and dates "
        "came in as plain text rather than a date type until explicitly converted with "
        "lubridate, both things python's pandas read_csv handled automatically with "
        "parse_dates.",
    ])

    doc.add_paragraph().add_run("C5.4 one open question").bold = True
    add_body(doc, [
        "The same open question from the python eda carries over unresolved here too: Figure "
        "5's correlation heatmap is a full 2006-2026 average, so we still do not know whether "
        "the chip cluster's tight correlation, or LVMH's low correlation with everyone, "
        "actually holds up specifically during 2008, 2020, and 2022, or whether it only looks "
        "stable because it is averaged over many calm years alongside the volatile ones. That "
        "is still the open thread from Assignment 3's H2 hypothesis work, not something this R "
        "pass answers on its own.",
    ])


def add_deliverables(doc):
    doc.add_page_break()
    add_section_heading(doc, "Deliverables")
    table = doc.add_table(rows=3, cols=1)
    table.cell(0, 0).text = f"Filename: {REG_NUMBER}_Team{TEAM_NUMBER}_Assignment4.docx"
    table.cell(1, 0).text = "R script: assignment4_eda.R (repo root, contains all code from Parts A and C)"
    table.cell(2, 0).text = "Submit via: Digiicampus"


def main():
    doc = Document()
    add_cover(doc)
    add_part_a(doc)
    add_part_b(doc)
    add_part_c0(doc)
    add_part_c1(doc)
    add_part_c2(doc)
    add_part_c3(doc)
    add_part_c4(doc)
    add_part_c5(doc)
    add_deliverables(doc)
    doc.save(OUTPUT_PATH)
    print(f"wrote {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
