"""Build the Assignment 2 Word deliverable from eda.ipynb.

Reads the executed notebook and rebuilds it as a docx: markdown becomes
formatted text, code cells become monospace blocks, and figure outputs
become embedded images. The interactive plotly figure has no static image
in the notebook, so its underlying chart is rebuilt here and exported to
PNG with kaleido.

Usage: uv run python scripts/build_docx.py
"""

import base64
import json
import re
from pathlib import Path

import pandas as pd
import plotly.express as px
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

REPO_ROOT = Path(__file__).resolve().parent.parent
NOTEBOOK_PATH = REPO_ROOT / "eda.ipynb"
OUTPUT_PATH = REPO_ROOT / "assignments" / "26120004_2_Assignment2.docx"

REG_NUMBER = "26120004"
TEAM_NUMBER = "2"
PROJECT_TITLE = "Historical Stock Analysis"
DATASET_LINE = "stock_data.csv, source: [add dataset link here]"
TEAM_MEMBERS = "[add team member names here]"
SUBMISSION_DATE = "2026-07-23"

MONO_FONT = "Consolas"


def add_cover_page(doc):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Assignment 2: Exploratory Data Analysis")
    run.bold = True
    run.font.size = Pt(20)

    doc.add_paragraph()

    fields = [
        ("Registration Number", REG_NUMBER),
        ("Team Number", TEAM_NUMBER),
        ("Project Title", PROJECT_TITLE),
        ("Dataset", DATASET_LINE),
        ("Team Members", TEAM_MEMBERS),
        ("Date", SUBMISSION_DATE),
    ]
    for label, value in fields:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(f"{label}: ")
        r.bold = True
        p.add_run(value)

    doc.add_page_break()


def add_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(13)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(10)
    p.add_run(text)


def add_paragraph_text(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(8)


def add_markdown_cell(doc, source):
    for block in source.split("\n\n"):
        b = block.strip()
        if not b:
            continue
        heading_match = re.fullmatch(r"\*\*(.+)\*\*", b)
        if heading_match:
            add_heading(doc, heading_match.group(1))
        elif b.startswith("- "):
            add_bullet(doc, b[2:].strip())
        else:
            add_paragraph_text(doc, b)


def add_output_text(doc, text):
    text = text.rstrip("\n")
    if not text:
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Inches(0.25)
    lines = text.split("\n")
    for i, line in enumerate(lines):
        run = p.add_run(line)
        run.font.name = MONO_FONT
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        if i < len(lines) - 1:
            run.add_break()


def add_output_image(doc, png_bytes):
    tmp_path = REPO_ROOT / "scripts" / "_tmp_fig.png"
    tmp_path.write_bytes(png_bytes)
    doc.add_picture(str(tmp_path), width=Inches(6))
    tmp_path.unlink()


def rebuild_plotly_figure_png():
    """Rebuild the Q2.3 interactive scatter and export a static PNG for the docx."""
    df = pd.read_csv(REPO_ROOT / "data" / "stock_data.csv", parse_dates=["Date"])
    df = df.sort_values(["Ticker", "Date"])
    df["Daily Return"] = df.groupby("Ticker")["Adj Close"].pct_change()
    plot_df = df.dropna(subset=["Daily Return"]).copy()
    plot_df["Abs Daily Return"] = plot_df["Daily Return"].abs()
    sample = plot_df.sample(3000, random_state=42).copy()
    sample["Date"] = sample["Date"].dt.strftime("%Y-%m-%d")

    fig = px.scatter(
        sample,
        x="Volume",
        y="Abs Daily Return",
        hover_data=["Ticker", "Date"],
        log_x=True,
        opacity=0.5,
        title="Figure 6: trading volume vs size of daily price move, interactive, 3,000 sampled ticker-days",
        labels={
            "Volume": "Volume, shares traded, log scale",
            "Abs Daily Return": "Absolute daily return (fraction)",
        },
    )
    return fig.to_image(format="png", width=1000, height=650, scale=2)


def get_output_text(output):
    if "text" in output:
        return "".join(output["text"])
    data = output.get("data", {})
    if "text/plain" in data:
        return "".join(data["text/plain"])
    return ""


def main():
    nb = json.loads(NOTEBOOK_PATH.read_text())
    plotly_png = None

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    add_cover_page(doc)

    for cell in nb["cells"]:
        source = "".join(cell["source"])
        if cell["cell_type"] == "markdown":
            add_markdown_cell(doc, source)
        elif cell["cell_type"] == "code":
            if not source.strip():
                continue
            # Q0.3 asks for "a table or a simple bar chart" of missing counts, so
            # that table is the deliverable itself, not incidental code output.
            is_missing_value_table = "isna().sum()" in source
            for output in cell.get("outputs", []):
                data = output.get("data", {})
                if "image/png" in data:
                    png_bytes = base64.b64decode(data["image/png"])
                    add_output_image(doc, png_bytes)
                elif "application/vnd.plotly.v1+json" in data:
                    if plotly_png is None:
                        plotly_png = rebuild_plotly_figure_png()
                    add_output_image(doc, plotly_png)
                elif is_missing_value_table:
                    add_output_text(doc, get_output_text(output))

    OUTPUT_PATH.parent.mkdir(exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"wrote {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
